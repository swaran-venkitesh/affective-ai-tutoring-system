from __future__ import annotations

import html
import json
import mimetypes
import os
import smtplib
import textwrap
from email.message import EmailMessage
from pathlib import Path
from typing import Any, Dict, Iterable, List

from local_secrets import load_local_env

load_local_env(Path(__file__).resolve().parent)


def _slugify(text: str) -> str:
    keep = "".join(ch if ch.isalnum() else "-" for ch in (text or "session").strip())
    keep = "-".join(part for part in keep.split("-") if part)
    return (keep or "session")[:64]


def _ensure_parent(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)


def _safe_float(value: Any, default: float = 0.0) -> float:
    try:
        return float(value)
    except Exception:
        return default


def _safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(value)
    except Exception:
        return default


def _graph_points(report: Dict[str, Any]) -> List[tuple[float, float, float]]:
    points: List[tuple[float, float, float]] = []
    for item in (report.get("engagement_timeline") or []):
        minute = _safe_float(item.get("t"), float(len(points)))
        score = _safe_float(item.get("score"), 0.0) * 100.0
        confusion = _safe_float(item.get("confusion"), 0.0) * 100.0
        points.append((minute, score, confusion))
    return points


def render_empathy_graph(report: Dict[str, Any], path: Path) -> Path | None:
    points = _graph_points(report)
    tone_points = list(report.get("tone_graph_timeline") or [])
    if not points and not tone_points:
        return None
    try:
        import matplotlib.pyplot as plt

        _ensure_parent(path)
        fig, ax = plt.subplots(figsize=(8.4, 3.6), dpi=160)
        if tone_points:
            xs = [float(item.get("t") or 0.0) for item in tone_points]
            text_scores = [None if item.get("text_score") is None else float(item.get("text_score")) for item in tone_points]
            camera_scores = [None if item.get("camera_score") is None else float(item.get("camera_score")) for item in tone_points]
            ax.plot(xs, text_scores, color="#4a90e2", linewidth=2.4, marker="o", markersize=3, label="Tutor text tone")
            if any(value is not None for value in camera_scores):
                ax.plot(xs, camera_scores, color="#cf4b3f", linewidth=2.2, marker="o", markersize=3, label="Camera monitoring")
            ax.set_title("Tutor Tone And Monitoring Timeline", fontsize=12, color="#5c3a21")
        else:
            xs = [p[0] for p in points]
            engagement = [p[1] for p in points]
            confusion = [p[2] for p in points]
            ax.plot(xs, engagement, color="#2e8b57", linewidth=2.4, marker="o", markersize=3, label="Engagement")
            ax.plot(xs, confusion, color="#c4483c", linewidth=2.0, linestyle="--", marker="o", markersize=3, label="Confusion/Stress")
            ax.set_title("Empathy Impact Through The Session", fontsize=12, color="#5c3a21")
        ax.set_facecolor("#f5ecd8")
        fig.patch.set_facecolor("#f1e3c2")
        ax.set_xlabel("Minutes", color="#5c3a21")
        ax.set_ylabel("Score", color="#5c3a21")
        ax.set_ylim(0, 100)
        ax.grid(alpha=0.22, color="#8d6b43")
        ax.tick_params(colors="#5c3a21")
        ax.legend(frameon=False, fontsize=9)
        for spine in ax.spines.values():
            spine.set_color("#8d6b43")
            spine.set_alpha(0.4)
        fig.tight_layout()
        fig.savefig(path, bbox_inches="tight")
        plt.close(fig)
        return path
    except Exception:
        return None


def build_material_text(report: Dict[str, Any]) -> str:
    sections: List[str] = []
    title = str(report.get("title") or "Session")
    topics = list(report.get("topics") or [])
    material_outline = str(report.get("material_outline") or "").strip()
    uploaded = list(report.get("uploaded_materials") or [])
    board_excerpt = str(report.get("board_material") or "").strip()

    sections.append(f"Session Material: {title}")
    if topics:
        sections.append("Topics covered:\n" + "\n".join(f"- {item}" for item in topics))
    if material_outline:
        sections.append("Study notes:\n" + material_outline)
    if board_excerpt:
        sections.append("Board notes:\n" + board_excerpt)
    if uploaded:
        preview = []
        for entry in uploaded[:8]:
            label = str(entry.get("label") or entry.get("filename") or "Uploaded file")
            content = str(entry.get("content") or "").strip()
            if content:
                preview.append(f"{label}:\n{content[:900]}")
        if preview:
            sections.append("Uploaded material context:\n" + "\n\n".join(preview))
    return "\n\n".join(section for section in sections if section.strip())


def _report_stats_rows(report: Dict[str, Any]) -> List[tuple[str, str]]:
    return [
        ("Student", str(report.get("student_name") or "Student")),
        ("Session ID", str(report.get("session_id") or "")),
        ("Mode", str(report.get("mode") or "shallow").title()),
        ("Date", str(report.get("date") or "")),
        ("Time", f"{report.get('start_time') or ''} - {report.get('end_time') or ''}".strip(" -")),
        ("Duration", f"{report.get('duration_min') or 0} min"),
        ("Turns", str(report.get("turn_count") or 0)),
        ("Engagement", f"{report.get('avg_engagement') or 0}%"),
        ("Attention", f"{report.get('attn_score') or 0}%"),
        ("Q&A Score", "N/A" if report.get("qa_score") is None else f"{report.get('qa_score')}%"),
    ]


def write_json_report(report: Dict[str, Any], path: Path) -> Path:
    _ensure_parent(path)
    path.write_text(json.dumps(report, indent=2), encoding="utf-8")
    return path


def write_png_report(report: Dict[str, Any], path: Path) -> Path:
    from PIL import Image, ImageDraw, ImageFont

    _ensure_parent(path)
    width = 1180
    padding = 56
    try:
        title_font = ImageFont.truetype("georgia.ttf", 42)
        heading_font = ImageFont.truetype("georgia.ttf", 28)
        body_font = ImageFont.truetype("georgia.ttf", 22)
        small_font = ImageFont.truetype("georgia.ttf", 19)
    except Exception:
        title_font = ImageFont.load_default()
        heading_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    insight_lines: List[str] = []
    for line in (report.get("empathy_summary") or []):
        insight_lines.extend(textwrap.wrap(str(line), width=82) or [str(line)])
    topic_lines = [f"- {t}" for t in (report.get("topics") or [])]
    tone_lines = []
    for item in (report.get("tone_timeline") or [])[:8]:
        tone = str(item.get("tone") or "supportive")
        cue = str(item.get("cue") or "")
        minute = item.get("t")
        label = f"{minute}m: {tone}" if minute is not None else tone
        if cue:
            label += f" - {cue}"
        tone_lines.extend(textwrap.wrap(label, width=84) or [label])

    stat_lines = [f"{label}: {value}" for label, value in _report_stats_rows(report)]
    all_lines = []
    all_lines.extend(stat_lines)
    all_lines.append("")
    all_lines.append("Empathy impact")
    all_lines.extend(insight_lines or ["No empathy summary available."])
    all_lines.append("")
    all_lines.append("Topics")
    all_lines.extend(topic_lines or ["- General"]) 
    if tone_lines:
        all_lines.append("")
        all_lines.append("Tutor tone timeline")
        all_lines.extend(tone_lines)

    line_height = 30
    height = max(900, padding * 2 + 180 + len(all_lines) * line_height)
    image = Image.new("RGB", (width, height), "#f1e3c2")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((18, 18, width - 18, height - 18), radius=28, fill="#f7edd4", outline="#8d6b43", width=3)
    draw.text((padding, padding), "Session Report", font=title_font, fill="#5c3a21")
    draw.text((padding, padding + 56), str(report.get("title") or "Session"), font=heading_font, fill="#7a5030")
    y = padding + 128
    draw.line((padding, y, width - padding, y), fill="#b54a3b", width=3)
    y += 28
    for idx, line in enumerate(all_lines):
        font = heading_font if line in {"Empathy impact", "Topics", "Tutor tone timeline"} else body_font
        fill = "#7a5030" if font is heading_font else "#3f2a18"
        draw.text((padding, y + idx * line_height), line, font=font, fill=fill)
    image.save(path)
    return path


def write_docx_report(report: Dict[str, Any], path: Path) -> Path:
    import docx

    _ensure_parent(path)
    doc = docx.Document()
    doc.add_heading(f"Session Report - {report.get('title') or 'Session'}", 0)
    stats = doc.add_table(rows=0, cols=2)
    for label, value in _report_stats_rows(report):
        row = stats.add_row().cells
        row[0].text = label
        row[1].text = value
    doc.add_heading("Topics Covered", level=1)
    for topic in (report.get("topics") or ["General"]):
        doc.add_paragraph(str(topic), style="List Bullet")
    doc.add_heading("Empathy Impact", level=1)
    for line in (report.get("empathy_summary") or ["No empathy summary available."]):
        doc.add_paragraph(str(line), style="List Bullet")
    doc.add_heading("Tutor Tone Timeline", level=1)
    for item in (report.get("tone_timeline") or [])[:10]:
        minute = item.get("t")
        tone = str(item.get("tone") or "supportive")
        cue = str(item.get("cue") or "")
        label = f"{minute}m" if minute is not None else "Turn"
        doc.add_paragraph(f"{label}: {tone} - {cue}".strip(" -"), style="List Bullet")
    doc.save(path)
    return path


def write_pdf_report(report: Dict[str, Any], path: Path, graph_path: Path | None = None) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Image, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    _ensure_parent(path)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleScroll", parent=styles["Title"], textColor=colors.HexColor("#5c3a21"), fontName="Helvetica-Bold")
    h_style = ParagraphStyle("HeadingScroll", parent=styles["Heading2"], textColor=colors.HexColor("#7a5030"), spaceAfter=8)
    body_style = ParagraphStyle("BodyScroll", parent=styles["BodyText"], textColor=colors.HexColor("#3f2a18"), leading=16)
    bullet_style = ParagraphStyle("BulletScroll", parent=styles["BodyText"], textColor=colors.HexColor("#3f2a18"), leftIndent=16, bulletIndent=4, leading=15)

    story = [
        Paragraph("Session Report", title_style),
        Paragraph(str(report.get("title") or "Session"), h_style),
        Spacer(1, 10),
    ]
    table = Table([[Paragraph(f"<b>{label}</b>", body_style), Paragraph(value, body_style)] for label, value in _report_stats_rows(report)], colWidths=[120, 340])
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#f7edd4")),
        ("BOX", (0, 0), (-1, -1), 1, colors.HexColor("#8d6b43")),
        ("INNERGRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#d1b891")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.extend([table, Spacer(1, 14), Paragraph("Topics Covered", h_style)])
    for topic in (report.get("topics") or ["General"]):
        story.append(Paragraph(str(topic), bullet_style, bulletText="-"))
    story.extend([Spacer(1, 10), Paragraph("Empathy Impact", h_style)])
    for line in (report.get("empathy_summary") or ["No empathy summary available."]):
        story.append(Paragraph(str(line), bullet_style, bulletText="-"))
    if graph_path and Path(graph_path).exists():
        story.extend([Spacer(1, 8), Image(str(graph_path), width=500, height=210)])
    tone_timeline = list(report.get("tone_timeline") or [])
    if tone_timeline:
        story.extend([Spacer(1, 10), Paragraph("Tutor Tone Timeline", h_style)])
        for item in tone_timeline[:10]:
            minute = item.get("t")
            tone = str(item.get("tone") or "supportive")
            cue = str(item.get("cue") or "")
            label = f"{minute}m" if minute is not None else "Turn"
            story.append(Paragraph(f"{label}: {tone} - {cue}".strip(" -"), bullet_style, bulletText="-"))
    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=36, rightMargin=36, topMargin=40, bottomMargin=36)
    doc.build(story)
    return path


def write_material_pdf(material_text: str, title: str, path: Path) -> Path:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    _ensure_parent(path)
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("MaterialTitle", parent=styles["Title"], textColor=colors.HexColor("#5c3a21"))
    body_style = ParagraphStyle("MaterialBody", parent=styles["BodyText"], textColor=colors.HexColor("#3f2a18"), leading=16)
    story = [Paragraph(f"Study Material - {title}", title_style), Spacer(1, 12)]
    for block in [b.strip() for b in (material_text or "").split("\n\n") if b.strip()]:
        for piece in textwrap.wrap(block, width=100, replace_whitespace=False):
            story.append(Paragraph(piece.replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 8))
    doc = SimpleDocTemplate(str(path), pagesize=A4, leftMargin=36, rightMargin=36, topMargin=40, bottomMargin=36)
    doc.build(story)
    return path


def generate_session_artifacts(report: Dict[str, Any], material_text: str, session_dir: Path) -> Dict[str, str]:
    slug = _slugify(f"{report.get('student_name') or 'student'}-{report.get('session_id') or 'session'}")
    report_dir = session_dir / "report"
    material_dir = session_dir / "material"
    graph_path = report_dir / f"{slug}_graph.png"
    graph_written = render_empathy_graph(report, graph_path)
    outputs = {
        "json": str(write_json_report(report, report_dir / f"{slug}.json")),
        "pdf": str(write_pdf_report(report, report_dir / f"{slug}.pdf", graph_written)),
        "docx": str(write_docx_report(report, report_dir / f"{slug}.docx")),
        "png": str(write_png_report(report, report_dir / f"{slug}.png")),
        "material_pdf": str(write_material_pdf(material_text, str(report.get("title") or "Session"), material_dir / f"{slug}_material.pdf")),
    }
    if graph_written:
        outputs["graph"] = str(graph_written)
    return outputs


def _smtp_config() -> Dict[str, Any]:
    host = os.getenv("TUTOR_SMTP_HOST", os.getenv("SMTP_HOST", "")).strip()
    port = _safe_int(os.getenv("TUTOR_SMTP_PORT", os.getenv("SMTP_PORT", "587")), 587)
    username = os.getenv("TUTOR_SMTP_USER", os.getenv("SMTP_USER", "")).strip()
    password = os.getenv("TUTOR_SMTP_PASS", os.getenv("SMTP_PASS", "")).strip()
    sender = os.getenv("TUTOR_SMTP_FROM", os.getenv("SMTP_FROM", username or "")).strip()
    use_tls = str(os.getenv("TUTOR_SMTP_TLS", "1")).strip().lower() not in {"0", "false", "no"}
    return {
        "host": host,
        "port": port,
        "username": username,
        "password": password,
        "sender": sender,
        "use_tls": use_tls,
    }


def _send_message(message: EmailMessage, config: Dict[str, Any]) -> Dict[str, Any]:
    if not config.get("host") or not config.get("sender"):
        return {"sent": False, "reason": "smtp_not_configured"}
    try:
        with smtplib.SMTP(str(config["host"]), int(config["port"]), timeout=25) as smtp:
            if config.get("use_tls"):
                smtp.starttls()
            if config.get("username"):
                smtp.login(str(config["username"]), str(config.get("password") or ""))
            smtp.send_message(message)
        return {"sent": True, "reason": "ok"}
    except Exception as exc:
        return {"sent": False, "reason": str(exc)}


def _escape_html(value: Any) -> str:
    return html.escape(str(value or ""))


def _wrap_email_html(title: str, subtitle: str, body_html: str, accent: str = "#b45309") -> str:
    safe_title = _escape_html(title)
    safe_subtitle = _escape_html(subtitle)
    return f"""\
<html>
  <body style="margin:0;padding:0;background:#f6efe2;font-family:Segoe UI,Arial,sans-serif;color:#2f2418;">
    <div style="max-width:640px;margin:0 auto;padding:24px 16px;">
      <div style="background:#fffaf0;border:1px solid #ead8b5;border-radius:18px;overflow:hidden;box-shadow:0 10px 30px rgba(92,58,33,0.10);">
        <div style="background:{accent};padding:18px 22px;color:#fff7ed;">
          <div style="font-size:26px;font-weight:700;line-height:1.2;">{safe_title}</div>
          <div style="margin-top:6px;font-size:14px;opacity:0.92;">{safe_subtitle}</div>
        </div>
        <div style="padding:24px 22px;line-height:1.7;font-size:15px;">
          {body_html}
        </div>
      </div>
    </div>
  </body>
</html>
"""


def send_session_email(recipient: str, report: Dict[str, Any], report_image_path: str, material_pdf_path: str) -> Dict[str, Any]:
    recipient = str(recipient or "").strip()
    if not recipient:
        return {"sent": False, "reason": "missing_recipient"}

    config = _smtp_config()
    if not config.get("host") or not config.get("sender"):
        return {"sent": False, "reason": "smtp_not_configured"}

    message = EmailMessage()
    message["Subject"] = "Your session report is ready 📚"
    message["From"] = str(config["sender"])
    message["To"] = recipient
    summary_lines = report.get("empathy_summary") or ["Your session artifacts are attached."]
    title = str(report.get("title") or "Session").strip() or "Session"
    student_name = str(report.get("student_name") or "Student").strip() or "Student"
    topics = [str(item).strip() for item in (report.get("topics") or []) if str(item).strip()]
    plain_lines = [
        f"Hello {student_name},",
        "",
        "Your session report and study material are ready 🎉",
        f"Session: {title}",
    ]
    if topics:
        plain_lines.append("Topics: " + ", ".join(topics[:5]))
    plain_lines.extend([""] + [str(line) for line in summary_lines[:4]])
    plain_lines.extend([
        "",
        "The report image and study material PDF are attached to this email 📎",
        "",
        "Regards,",
        "Tutoring System",
    ])
    message.set_content("\n".join(plain_lines))
    topic_html = ""
    if topics:
        topic_html = "<ul style=\"margin:10px 0 14px 20px;padding:0;\">" + "".join(
            f"<li style=\"margin:4px 0;\">{_escape_html(item)}</li>" for item in topics[:6]
        ) + "</ul>"
    summary_html = "".join(
        f"<li style=\"margin:6px 0;\">{_escape_html(line)}</li>" for line in summary_lines[:5]
    )
    body_html = (
        f"<p style=\"margin:0 0 12px;\">Hello <strong>{_escape_html(student_name)}</strong> 👋</p>"
        f"<p style=\"margin:0 0 12px;\">Your session report and study material are ready 🎉</p>"
        f"<div style=\"background:#fff3d6;border:1px solid #efd6a2;border-radius:14px;padding:14px 16px;margin:14px 0;\">"
        f"<div style=\"font-size:13px;color:#8a5b22;text-transform:uppercase;letter-spacing:0.05em;\">Session</div>"
        f"<div style=\"font-size:20px;font-weight:700;color:#5c3a21;\">{_escape_html(title)}</div>"
        f"</div>"
        f"{topic_html}"
        f"<p style=\"margin:0 0 8px;\">Highlights from this session ✨</p>"
        f"<ul style=\"margin:8px 0 14px 20px;padding:0;\">{summary_html}</ul>"
        f"<p style=\"margin:0;\">The report image and study material PDF are attached to this email 📎</p>"
    )
    message.add_alternative(
        _wrap_email_html("Session Complete 🎓", "Your learning artifacts are attached", body_html, accent="#a16207"),
        subtype="html",
    )

    for file_path in [report_image_path, material_pdf_path]:
        if not file_path:
            continue
        path = Path(file_path)
        if not path.exists():
            continue
        mime, _ = mimetypes.guess_type(str(path))
        maintype, subtype = (mime.split("/", 1) if mime else ("application", "octet-stream"))
        message.add_attachment(path.read_bytes(), maintype=maintype, subtype=subtype, filename=path.name)
    return _send_message(message, config)


def send_welcome_email(recipient: str, student_name: str, student_id: str = "") -> Dict[str, Any]:
    recipient = str(recipient or "").strip()
    if not recipient:
        return {"sent": False, "reason": "missing_recipient"}

    config = _smtp_config()
    if not config.get("host") or not config.get("sender"):
        return {"sent": False, "reason": "smtp_not_configured"}

    name = str(student_name or "Student").strip() or "Student"
    sid = str(student_id or "").strip()
    message = EmailMessage()
    message["Subject"] = "Welcome to the Tutoring System 🎉"
    message["From"] = str(config["sender"])
    message["To"] = recipient
    lines = [
        f"Hello {name},",
        "",
        "Your tutoring account has been created successfully 🎉",
        "You can now start guided sessions, get reports, and receive study materials by email.",
    ]
    if sid:
        lines.extend(["", f"Student ID: {sid}"])
    lines.extend([
        "",
        "You will receive session report and material emails after end-session generation when your session includes email delivery 📩",
        "",
        "Regards,",
        "Tutoring System",
    ])
    message.set_content("\n".join(lines))
    id_html = ""
    if sid:
        id_html = (
            f"<div style=\"display:inline-block;background:#fff3d6;border:1px solid #efd6a2;border-radius:999px;"
            f"padding:8px 14px;margin:6px 0 16px;color:#6b4226;font-weight:600;\">"
            f"Student ID: {_escape_html(sid)}</div>"
        )
    body_html = (
        f"<p style=\"margin:0 0 12px;\">Hello <strong>{_escape_html(name)}</strong> 👋</p>"
        f"<p style=\"margin:0 0 12px;\">Your tutoring account has been created successfully 🎉</p>"
        f"{id_html}"
        f"<p style=\"margin:0 0 10px;\">You are ready to start guided learning sessions, receive useful reports, and get study material by email 📚</p>"
        f"<ul style=\"margin:8px 0 14px 20px;padding:0;\">"
        f"<li style=\"margin:6px 0;\">Start a fresh tutoring session whenever you need.</li>"
        f"<li style=\"margin:6px 0;\">Receive session report and material emails after end session.</li>"
        f"<li style=\"margin:6px 0;\">Continue with your saved student profile and preferences.</li>"
        f"</ul>"
        f"<p style=\"margin:0;\">Keep learning and keep going 🚀</p>"
    )
    message.add_alternative(
        _wrap_email_html("Account Created 🎊", "Your tutoring access is ready", body_html, accent="#b45309"),
        subtype="html",
    )
    return _send_message(message, config)
